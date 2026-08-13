import os
import re
import json
import requests
import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Pt, Mm  

def extract_code_blocks(py_file_path):
    """提取源代码块（兼容宽容正则匹配）"""
    code_blocks = {}
    current_block = None
    current_content = []
    
    start_pattern = re.compile(r'#\s*[-]+\s*([a-zA-Z0-9_]+)\s*[-]+')
    end_pattern = re.compile(r'#\s*[-]+\s*end\s*[-]+')

    if not os.path.exists(py_file_path):
        return code_blocks

    with open(py_file_path, 'r', encoding='utf-8') as f:
        for line in f:
            start_match = start_pattern.search(line)
            end_match = end_pattern.search(line)

            if start_match and not end_match:
                current_block = start_match.group(1)
                current_content = []
            elif end_match and current_block:
                code_blocks[current_block] = "".join(current_content).strip()
                current_block = None
            elif current_block is not None:
                current_content.append(line)
    return code_blocks

def call_ai_for_text(code_text):
    """
    真正的 Agent 大脑：通过本地 Ollama 引擎，让大模型分析代码并返回文案
    """
    print("🧠 正在唤醒本地大模型阅读代码并撰写文案...")
    
    # 1. 设定 Agent 的系统提示词 (Prompt)
    prompt = f"""
    你是一个严谨的计算机视觉实验报告生成助理。
    请阅读以下 Python 实验代码，提取核心技术并为其撰写实验文案。
    
    代码内容：
    ```python
    {code_text}
    ```
    
    请严格按照以下 JSON 格式返回，禁止输出任何 Markdown 代码块标记及多余废话：
    1. "exp_title": 根据代码提炼一个专业的实验大标题，例如“基于YOLOv5的机器视觉目标检测实验”。
    2. "exp_purpose": 使用数字序号（1. 2. 3.）分点作答。每一条必须是“掌握.../学习.../了解...”开头的完整句子。换行请使用 \\n。
    3. "exp_content": 用一段 50 到 80 字的连贯文字描述实验的具体内容。
    4. "exp_theory": 【核心要求】请撰写约150字的实验原理，必须包含学术名词。格式为：“1.核心算法原理：... \\n2.推理处理流程：... \\n3.实际应用场景：...”。绝对禁止只输出数字！
    5. "step_1": 固定输出：“1. 运行jupyter lab\\n(1)打开桌面的“实验”文件夹，在终端输入 jupyter lab；\\n(2)新建 Notebook 并命名为实验名称。”
    6. "step_2": 提炼导入了哪些库，如：“2. 导入库文件及函数\\n导入...等相关库。”
    7. "step_3": 提炼加载了什么模型，如：“3. 加载模型\\n设置路径并初始化...模型。”
    8. "step_4": 概括主程序流程，如：“4. 主程序\\n打开相机获取视频流，进行循环推理检测...”
    
    {{
        "exp_title": "...",
        "exp_purpose": "1. 掌握...\\n2. 学习...",
        "exp_content": "...",
        "exp_theory": "1.核心算法原理：YOLOv5是一种...\\n2.推理处理流程：首先读取相机...\\n3.实际应用场景：智能监控...",
        "step_1": "...",
        "step_2": "...",
        "step_3": "...",
        "step_4": "..."
    }}
    """

    # 2. 调用本地 Ollama API
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "qwen2.5:3b",  
        "prompt": prompt,
        "format": "json",       
        "stream": False
    }

    try:
        response = requests.post(url, json=payload)
        response.raise_for_status()
        result_text = response.json().get("response", "")
        
        # 3. 解析 AI 返回的 JSON 数据
        ai_response = json.loads(result_text)
        print("💡 AI 文案撰写完成！")
        return ai_response
    except Exception as e:
        print(f"❌ 调用本地 AI 失败，请检查 Ollama 是否运行: {e}")
        return {"exp_purpose": "AI生成失败", "exp_content": "AI生成失败", "exp_theory": "AI生成失败"}

def build_report(template_path, py_file, img_file, output_path):
    print("1. 读取并解析 Python 源代码...")
    code_blocks = extract_code_blocks(py_file)
    
    with open(py_file, 'r', encoding='utf-8') as f:
        full_code_text = f.read()
    
    print("2. 召唤 AI 撰写实验文案...")
    text_data = call_ai_for_text(full_code_text)

    print("3. 加载 Word 模板并装配图片...")
    try:
        doc = DocxTemplate(template_path)
    except Exception as e:
        print(f"模板加载失败: {e}")
        return
        
    # 安全且正确地处理图片：限制宽度为 140 毫米
    if os.path.exists(img_file):
        image_obj = InlineImage(doc, img_file, width=Mm(140))
    else:
        image_obj = "【提示：未找到运行结果截图】"

    # 组装最终的数据字典
    context = {
        # 👇 新增这一行，接收大标题
        "exp_title": text_data.get("exp_title", "计算机视觉自动化实验报告"),
        
        "exp_purpose": text_data.get("exp_purpose", ""),
        "exp_content": text_data.get("exp_content", ""),
        "exp_theory": text_data.get("exp_theory", ""),
        "step_1": text_data.get("step_1", "1. 运行环境设置"),
        "step_2": text_data.get("step_2", "2. 导入库文件"),
        "step_3": text_data.get("step_3", "3. 加载模型"),
        "step_4": text_data.get("step_4", "4. 主程序实现"),
        "code_import": code_blocks.get("code_import", "缺失导入代码"),
        "code_load_model": code_blocks.get("code_load_model", "缺失加载代码"),
        "code_main": code_blocks.get("code_main", "缺失主程序代码"),
        "result_image": image_obj
    }

    print("4. 正在渲染生成 Word 文档...")
    doc.render(context)
    doc.save(output_path) # 先保存一次，让模板和图片生效

    print("5. 正在为代码块绘制专业边框...")
    # 重新加载刚才生成的文档进行二次加工
    final_doc = docx.Document(output_path)
    
    # 遍历所有段落，寻找“参考代码：”
    for i in range(len(final_doc.paragraphs)):
        if "参考代码：" in final_doc.paragraphs[i].text:
            # 找到紧挨着的下一段（就是我们填入的代码段）
            if i + 1 < len(final_doc.paragraphs):
                code_para = final_doc.paragraphs[i+1]

                # 强制撑开段前和段后距离，防止边框撞到字
                code_para.paragraph_format.space_before = Pt(12) 
                code_para.paragraph_format.space_after = Pt(12)

                # 绘制边框的底层 XML 魔法
                pPr = code_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')    
                    border.set(qn('w:sz'), '8')          
                    border.set(qn('w:space'), '6')       
                    border.set(qn('w:color'), '808080')  
                    pBdr.append(border)
                pPr.append(pBdr)
                
    # 覆盖保存最终完美版
    final_doc.save(output_path)
    print(f"✅ 大功告成！终极版报告已保存至：{output_path}")

if __name__ == "__main__":
    TEMPLATE = "template.docx"          
    SOURCE_CODE = "experiment.py"       
    RESULT_IMG = "result.jpg"           
    OUTPUT_REPORT = "AI自动生成_实验文档.docx"

    build_report(TEMPLATE, SOURCE_CODE, RESULT_IMG, OUTPUT_REPORT)